def preprocess_text(text):
    import re
    # Replace more than two newline characters with exactly two newline characters
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text

def get_printable_line(line):
    import re
    # Define substitutions for specific non-printable characters
    # This is just an example; you can expand it based on your needs
    substitutions = {
        '\t': '[TAB]',
        'É' : "E",
        'Í' : 'I',
        'º' : '*',
        '' : '---',
    }

    # Replace specific control characters
    for key, value in substitutions.items():
        line = line.replace(key, value)

    # Replace other non-printable characters with a placeholder
    # This regex matches any character that is not printable ASCII, excluding newline
    line = re.sub(r'[^\x20-\x7E\n]', '*', line)
    return line

def output_to_word(output_path: str, remove_origin=True):
    from docx import Document
    from docx.shared import Pt  # For font size
    import os

    # Create a new Document
    doc = Document()

    # Open the text file and read lines
    with open(output_path, 'r',  encoding='ISO-8859-1') as file:
        text_content = file.read()

    # Preprocess the entire text content to merge excessive line breaks
    preprocessed_text = preprocess_text(text_content)

    # Split the preprocessed text into lines
    lines = preprocessed_text.split('\n')

    # Add lines to the Word document
    for line in lines:
        try:
            paragraph = doc.add_paragraph(line)
        except:
            paragraph = doc.add_paragraph(get_printable_line(line))
        # Set font size for this paragraph
        for run in paragraph.runs:
            run.font.size = Pt(8)  # Set font size

    docx_file_path = output_path.split(".")[0] + ".docx"
    # Save the document
    doc.save(docx_file_path)
    if remove_origin:
        os.remove(output_path)
